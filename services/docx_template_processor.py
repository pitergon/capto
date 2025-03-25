from docx import Document
import copy
import re
from utils.logger import botlog


class DocxTemplateProcessor:
    def __init__(self, template_path: str, output_path: str, replacements: dict, placeholder_pattern: str = r'{{([a-zA-Z0-9._+-]+)}}'):
        """Initializes the template processor with file paths and replacements dictionary."""
        self.template_path = template_path
        self.output_path = output_path
        self.replacements = replacements
        self.placeholder_pattern = placeholder_pattern
        self.doc = Document(template_path)

    def _find_header_by_style(self, start_idx=0, backward=False):
        """Finds the nearest header based on style."""
        return self._find_prev_header_by_style(start_idx) if backward else self._find_next_header_by_style(start_idx)

    def _find_next_header_by_style(self, start_idx=0):
        """Finds the next header based on style."""
        for idx, p in enumerate(self.doc.paragraphs[start_idx:], start=start_idx):
            if p.style.name.startswith("Heading"):
                return idx, p

    def _find_prev_header_by_style(self, start_idx=0):
        """Finds the previous header based on style."""
        for offset, p in enumerate(self.doc.paragraphs[start_idx::-1]):
            if p.style.name.startswith("Heading"):
                return start_idx - offset, p

    def _remove_all_placeholders(self, text):
        """Removes all placeholders from a string."""
        return re.sub(self.placeholder_pattern, "", text)

    def _select_block(self, idx):
        """Selects a block of paragraphs around the given paragraph index, excluding headers."""
        block = []
        prev_paragraphs = self.doc.paragraphs[idx::-1]
        next_paragraphs = self.doc.paragraphs[idx + 1:]
        for p in prev_paragraphs:
            if p.style.name.startswith("Heading"):
                break
            block.append(p)
        block.reverse()
        for p in next_paragraphs:
            if p.style.name.startswith("Heading"):
                break
            block.append(p)
        return block

    def _get_empty_paragraph(self):
        """Creates and returns an empty paragraph to serve as a separator."""
        new_p = self.doc.add_paragraph("")
        empty_p = copy.deepcopy(new_p)
        new_p._element.getparent().remove(new_p._element)
        return empty_p

    def _insert_block_after(self, paragraphs_list, paragraph=None, idx=None):
        """Inserts a list of paragraphs after a given paragraph."""
        if not paragraph and idx is not None and idx >= 0:
            paragraph = self.doc.paragraphs[idx] if idx < len(self.doc.paragraphs) else self.doc.paragraphs[-1]
        for p in reversed(paragraphs_list):
            paragraph._element.addnext(p._element)

    def _find_placeholders_in_block(self, paragraph_list):
        """Finds all placeholders in a block of paragraphs."""
        placeholders = []
        for p in paragraph_list:
            matches = re.findall(self.placeholder_pattern, p.text)
            if matches:
                placeholders.append(*matches)
        return placeholders

    def _get_placeholder_value(self, match, multi_value_idx=None):
        """
            Returns the replacement value for a given placeholder.
            - If the key is not found in the replacements dictionary, returns an empty string.
            - If the value in the dictionary is a list, returns the value at index `multi_value_idx`.
            - If `multi_value_idx` is out of range for the list, returns an empty string.
            - If the value is not a list, returns it as is.
            - If the value is a list and `multi_value_idx` is None, returns the original placeholder without replacement for next processing
        """
        key = match.group(1)
        value = self.replacements.get(key, "")
        if isinstance(value, list) and multi_value_idx is not None:
            return value[multi_value_idx] if multi_value_idx < len(value) else ""
        return value if not isinstance(value, list) else match.group(0)

    def _process_paragraph(self, paragraph, multi_value_idx=None):
        """
        Processes a paragraph, replacing placeholders with their values.
        - To preserve formatting within a single paragraph, ensure that the placeholder is entered as a whole
          or inserted using the "Paste as plain text" option.
        - If a placeholder is split across multiple runs, it will still be replaced, but the paragraph's formatting will be lost.
        """
        for run in paragraph.runs:
            if re.search(self.placeholder_pattern, run.text):
                run.text = re.sub(self.placeholder_pattern, lambda match: self._get_placeholder_value(match, multi_value_idx), run.text)

        placeholders = re.findall(self.placeholder_pattern, paragraph.text)
        contains_single_value_ph = any(
            key not in self.replacements or not isinstance(self.replacements[key], list) for key in placeholders
        )

        if placeholders:
            if multi_value_idx is not None or contains_single_value_ph:
                paragraph.text = re.sub(self.placeholder_pattern, lambda match: self._get_placeholder_value(match, multi_value_idx), paragraph.text)

    def _process_multi_value_placeholders(self, p_idx):
        """Handles multi-value placeholders by duplicating the paragraph block and replacing values in copies."""
        selected_block = self._select_block(p_idx)
        placeholders = self._find_placeholders_in_block(selected_block)
        if not placeholders:
            return
        copies_count = max(
            (len(self.replacements[key]) for key in placeholders if key in self.replacements and isinstance(self.replacements[key], list)),
            default=0
        )
        blocks_list = [selected_block] + [copy.deepcopy(selected_block) for _ in range(copies_count - 1)]
        paragraphs_list = []

        for block_num, block in enumerate(blocks_list):
            for p in block:
                self._process_paragraph(p, multi_value_idx=block_num)
                if block_num > 0:
                    paragraphs_list.append(p)
            paragraphs_list.append(self._get_empty_paragraph())

        last_paragraph = selected_block[-1]
        self._insert_block_after(paragraphs_list, last_paragraph)

    def process_document(self):
        """Processes the entire document, replacing placeholders and handling multi-value placeholders."""
        multi_value_paragraphs_idx = []
        for idx, paragraph in enumerate(self.doc.paragraphs):
            placeholders = re.findall(self.placeholder_pattern, paragraph.text)
            contains_multi_value_ph = any(isinstance(self.replacements.get(key), list) for key in placeholders)
            if contains_multi_value_ph:
                multi_value_paragraphs_idx.append((idx, paragraph.text))  # text saved for debug
            self._process_paragraph(paragraph)

        paragraph_count = len(self.doc.paragraphs)
        for idx, _ in multi_value_paragraphs_idx:
            offset = len(self.doc.paragraphs) - paragraph_count  # Adjust index after insertions
            self._process_multi_value_placeholders(idx + offset)

    def save(self):
        """Saves the processed document to the specified output path."""
        self.doc.save(self.output_path)
        botlog(f"File saved to {self.output_path}")
